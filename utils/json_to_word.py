# -*- coding: utf-8 -*-
"""
将“政策干预仿真”JSON结果转换为格式化 Word（.docx）报告。
"""
from __future__ import annotations
import argparse
import json
from typing import Any, Dict, List
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

def set_style_font(style, name: str = "微软雅黑", size_pt: float = 11, bold: bool = False) -> None:
    """字体设置"""
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    font.bold = bold
    rFonts = style._element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), name)

def configure_document(doc: Document) -> None:
    """基础页面排版设置"""
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    set_style_font(doc.styles["Normal"], "微软雅黑", 11, False)
    if "Heading 1" in doc.styles:
        set_style_font(doc.styles["Heading 1"], "微软雅黑", 16, True)
    if "Heading 2" in doc.styles:
        set_style_font(doc.styles["Heading 2"], "微软雅黑", 13, True)
    if "Heading 3" in doc.styles:
        set_style_font(doc.styles["Heading 3"], "微软雅黑", 12, True)
    if "List Bullet" in doc.styles:
        set_style_font(doc.styles["List Bullet"], "微软雅黑", 11, False)

def format_range(rng: Any) -> str:
    """Format [min,max] -> 'min–max'."""
    if rng is None:
        return "未提供"
    if isinstance(rng, (list, tuple)) and len(rng) == 2:
        a, b = rng
        def fmt(x: Any) -> str:
            if x is None:
                return "null"
            if isinstance(x, int):
                return str(x)
            if isinstance(x, float):
                if abs(x - round(x)) < 1e-9:
                    return str(int(round(x)))
                s = f"{x:.2f}".rstrip("0").rstrip(".")
                return s
            return str(x)

        return f"{fmt(a)}–{fmt(b)}"
    return str(rng)

def add_label_paragraph(doc: Document, label: str, text: str = "") -> None:
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    if text:
        p.add_run(text)

def add_bullets(doc: Document, items: List[str]) -> None:
    for it in items or []:
        doc.add_paragraph(str(it), style="List Bullet")

def add_evidence_inline(doc: Document, evidence: List[Dict[str, Any]]) -> None:
    if not evidence:
        add_label_paragraph(doc, "证据引用：", "无")
        return
    parts = []
    for e in evidence:
        cid = e.get("chunk_id")
        did = e.get("doc_id")
        if cid or did:
            parts.append(f"{cid or ''}/{did or ''}".strip("/"))
    add_label_paragraph(doc, "证据引用：", "；".join(parts) if parts else "无")


def get_pol_section(pol: Dict[str, Any], x: Any, sec: str) -> Dict[str, Any]:
    """
    返回第三章政策仿真字典名 3.x.1~3.x.4.
    兼容两种JSON键名：
      - "3.x.1"（字面量x，当前你输出的就是这种）
      - "3.<x>.1"（把x展开成具体编号，如3.1.1）
    """
    key_literal = f"3.x.{sec}"
    if key_literal in pol and isinstance(pol[key_literal], dict):
        return pol[key_literal]

    key_expanded = f"3.{x}.{sec}"
    if key_expanded in pol and isinstance(pol[key_expanded], dict):
        return pol[key_expanded]

    # 匹配任何以"3."开头且以".<sec>"结尾的键
    for k, v in pol.items():
        if isinstance(k, str) and k.startswith("3.") and k.endswith(f".{sec}") and isinstance(v, dict):
            return v

    return {}


def json_report_to_docx(data: Dict[str, Any], out_path: str) -> str:
    doc = Document()
    configure_document(doc)

    # 第一章 引言
    ch1 = data.get("chapter1", {})
    doc.add_heading("第一章 引言", level=1)
    sec11 = ch1.get("1.1", {})
    doc.add_heading("1.1 新能源汽车行业内卷现象概述", level=2)

    if sec11.get("text"):
        doc.add_paragraph(sec11["text"])

    bullets = sec11.get("bullets") or []
    if bullets:
        add_label_paragraph(doc, "主要表现（要点）：")
        add_bullets(doc, bullets)

    risks = sec11.get("key_risks") or []
    if risks:
        add_label_paragraph(doc, "关键风险：")
        add_bullets(doc, risks)

    add_evidence_inline(doc, sec11.get("evidence") or [])
    doc.add_paragraph()

    # 第二章 行业状态
    ch2 = data.get("chapter2", {})
    doc.add_heading("第二章 行业状态", level=1)

    sec21 = ch2.get("2.1", {})
    doc.add_heading("2.1 当前新能源行业状态", level=2)
    if sec21.get("text"):
        doc.add_paragraph(sec21["text"])
    if sec21.get("bullets"):
        add_label_paragraph(doc, "状态要点：")
        add_bullets(doc, sec21["bullets"])
    if "involution_index_baseline_range" in sec21:
        add_label_paragraph(doc, "内卷指数（基准）：", format_range(sec21.get("involution_index_baseline_range")))
    #if sec21.get("confidence") is not None:
    #    add_label_paragraph(doc, "置信度：", f"{float(sec21.get('confidence')):.2f}")
    add_evidence_inline(doc, sec21.get("evidence") or [])
    doc.add_paragraph()

    sec22 = ch2.get("2.2", {})
    doc.add_heading("2.2 未来趋势预测", level=2)
    if sec22.get("text"):
        doc.add_paragraph(sec22["text"])
    if sec22.get("bullets"):
        add_label_paragraph(doc, "趋势要点：")
        add_bullets(doc, sec22["bullets"])
    if sec22.get("trend_points"):
        add_label_paragraph(doc, "关键趋势点：")
        add_bullets(doc, sec22["trend_points"])
    if sec22.get("risk_triggers"):
        add_label_paragraph(doc, "风险触发条件：")
        add_bullets(doc, sec22["risk_triggers"])
    add_evidence_inline(doc, sec22.get("evidence") or [])
    doc.add_paragraph()

    # 第三章 政策情景设定
    ch3 = data.get("chapter3", {})
    doc.add_heading("第三章 政策情景设定", level=1)

    policies = ch3.get("policies") or []
    for pol in policies:
        x = pol.get("x")
        name = pol.get("policy_name") or f"政策{x}"
        doc.add_heading(f"政策 {x}：{name}", level=2)

        # 3.x.1 政策内容
        s1 = get_pol_section(pol, x, "1")
        doc.add_heading(f"3.{x}.1 政策内容", level=3)

        measures = s1.get("policy_measures") or []
        if measures:
            add_label_paragraph(doc, "政策措施：")
            add_bullets(doc, measures)

        params = s1.get("parameters") or []
        if params:
            add_label_paragraph(doc, "参数：")
            table = doc.add_table(rows=1, cols=3)
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            hdr = table.rows[0].cells
            hdr[0].text = "参数名"
            hdr[1].text = "取值"
            hdr[2].text = "说明"
            for prm in params:
                row = table.add_row().cells
                row[0].text = str(prm.get("name") or "")
                row[1].text = "" if prm.get("value") is None else str(prm.get("value"))
                row[2].text = str(prm.get("note") or "")
        doc.add_paragraph()

        # 3.x.2 政策作用机制
        s2 = get_pol_section(pol, x, "2")
        doc.add_heading(f"3.{x}.2 政策作用机制", level=3)
        chain = s2.get("mechanism_chain") or []
        if chain:
            add_label_paragraph(doc, "作用链条：")
            add_bullets(doc, chain)
        levers = s2.get("primary_levers") or []
        if levers:
            add_label_paragraph(doc, "主要作用杠杆：", "、".join(levers))
        doc.add_paragraph()

        # 3.x.3 适用场景与边界条件
        s3 = get_pol_section(pol, x, "3")
        doc.add_heading(f"3.{x}.3 适用场景与边界条件", level=3)
        if s3.get("applicable_when"):
            add_label_paragraph(doc, "适用场景：")
            add_bullets(doc, s3["applicable_when"])
        if s3.get("boundary_conditions"):
            add_label_paragraph(doc, "边界条件：")
            add_bullets(doc, s3["boundary_conditions"])
        if s3.get("failure_modes"):
            add_label_paragraph(doc, "失效模式：")
            add_bullets(doc, s3["failure_modes"])
        doc.add_paragraph()

        # 3.x.4 政策对企业行为/产业行为的影响
        s4 = get_pol_section(pol, x, "4")
        doc.add_heading(f"3.{x}.4 政策对企业行为/产业行为的影响", level=3)

        inv = s4.get("involution_index") or {}
        if inv:
            add_label_paragraph(doc, "内卷指数影响：")
            add_bullets(
                doc,
                [
                    f"基准范围：{format_range(inv.get('baseline_range'))}",
                    f"政策后范围：{format_range(inv.get('after_range'))}",
                    f"变化范围：{format_range(inv.get('change_range'))}",
                    #f"置信度：{float(inv.get('confidence')):.2f}" if inv.get("confidence") is not None else "置信度：未提供",
                ],
            )

        beh = s4.get("behavior_impacts") or {}
        if beh:
            add_label_paragraph(doc, "行为影响：")
            order = [
                ("pricing", "定价"),
                ("capacity", "产能"),
                ("rnd", "研发"),
                ("channels", "渠道"),
                ("supply_chain_terms", "供应链账期"),
                ("mna_exit", "并购退出"),
            ]
            for k, cn in order:
                item = beh.get(k) or {}
                direction = item.get("direction") or "mixed"
                txt = item.get("text") or ""
                doc.add_paragraph(f"{cn}（{direction}）：{txt}", style="List Bullet")

        if s4.get("kpis"):
            add_label_paragraph(doc, "关键 KPI：")
            add_bullets(doc, s4["kpis"])
        if s4.get("side_effects"):
            add_label_paragraph(doc, "潜在副作用：")
            add_bullets(doc, s4["side_effects"])

        doc.add_paragraph()

    # 第四章 政策建议
    ch4 = data.get("chapter4", {})
    doc.add_heading("第四章 政策建议", level=1)

    sec41 = ch4.get("4.1", {})
    doc.add_heading("4.1 推荐方案（主推/备选/不建议）", level=2)
    for label, key in [("主推", "primary"), ("备选", "secondary"), ("不建议", "not_recommended")]:
        items = sec41.get(key) or []
        if items:
            add_label_paragraph(doc, f"{label}：")
            for it in items:
                doc.add_paragraph(f"政策 {it.get('policy_x')}：{it.get('why')}", style="List Bullet")
    if not (sec41.get("primary") or sec41.get("secondary") or sec41.get("not_recommended")):
        doc.add_paragraph("（无）")
    doc.add_paragraph()

    sec42 = ch4.get("4.2", {})
    doc.add_heading("4.2 分场景选择规则", level=2)
    rules = sec42.get("rules") or []
    if not rules:
        doc.add_paragraph("（无）")
    else:
        for rule in rules:
            scene = rule.get("scene") or "未命名场景"
            doc.add_heading(f"场景：{scene}", level=3)
            if rule.get("triggers"):
                add_label_paragraph(doc, "触发条件：")
                add_bullets(doc, rule["triggers"])
            if rule.get("recommended_policy_x"):
                add_label_paragraph(doc, "推荐政策：", "、".join([str(x) for x in rule["recommended_policy_x"]]))
            if rule.get("expected_results"):
                add_label_paragraph(doc, "预期结果：")
                add_bullets(doc, rule["expected_results"])
            if rule.get("watchouts"):
                add_label_paragraph(doc, "注意事项：")
                add_bullets(doc, rule["watchouts"])
            doc.add_paragraph()

    sec43 = ch4.get("4.3", {})
    doc.add_heading("4.3 配套机制及注意事项（监管、信息披露、退出与消费者保障等）", level=2)
    for label, key in [
        ("配套机制", "supporting_mechanisms"),
        ("监管、信息披露", "governance_and_disclosure"),
        ("退出与消费者保障", "exit_and_consumer_protection"),
        ("监测评估与迭代", "monitoring_and_iteration"),
    ]:
        items = sec43.get(key) or []
        if items:
            add_label_paragraph(doc, f"{label}：")
            add_bullets(doc, items)

    # 备注与免责声明
    #notes = data.get("notes") or []
    #disclaimer = data.get("disclaimer")
    #if notes or disclaimer:
    #    doc.add_paragraph()
    #    doc.add_heading("备注与免责声明", level=1)
    #    if notes:
    #        add_label_paragraph(doc, "备注：")
    #        add_bullets(doc, notes)
    #    if disclaimer:
    #        add_label_paragraph(doc, "免责声明：", str(disclaimer))

    doc.save(out_path)
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--json", default="output/policy_sim_NEV_20260109_095559.json", help="输入 JSON 文件路径")
    parser.add_argument("--out", default="output/policy_word.docx", help="输出 DOCX 文件路径")
    args = parser.parse_args()
    with open(args.json, "r", encoding="utf-8") as f:
        data = json.load(f)
    json_report_to_docx(data, args.out)
    print(f"已生成：{args.out}")

# ---------------------------
# 本地快速测试
# ---------------------------
if __name__ == "__main__":
    main()
